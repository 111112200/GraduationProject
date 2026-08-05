import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.services.docx_parser_service import parse_docx_report


class DocxParserRegressionTest(unittest.TestCase):
    def _parse(self, build_document):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "report.docx"
            document = Document()
            build_document(document)
            document.save(path)
            return parse_docx_report(str(path))

    def test_body_keywords_do_not_change_section_state(self):
        def build(document):
            document.add_paragraph("引言正文提到心得体会，但这一整句不是章节标题。")
            document.add_paragraph("设计思路", style="Heading 2")
            document.add_paragraph("采用分层结构完成实验模块设计。")
            document.add_paragraph("4.2 实验结果", style="Heading 2")
            document.add_paragraph("结果分析不应继续归入设计思路。")

        blocks = self._parse(build)

        self.assertEqual([block["section_type"] for block in blocks], ["DESIGN_IDEA"])
        self.assertEqual(blocks[0]["content"], "采用分层结构完成实验模块设计。")
        self.assertFalse(blocks[0]["fallback"])

    def test_inline_heading_content_is_preserved(self):
        def build(document):
            document.add_paragraph("心得体会：这次实验让我掌握了排序算法。")
            document.add_paragraph("后续反思内容。")

        blocks = self._parse(build)

        self.assertEqual(len(blocks), 2)
        self.assertEqual(blocks[0]["section_type"], "REFLECTION")
        self.assertEqual(blocks[0]["content"], "这次实验让我掌握了排序算法。")
        self.assertEqual(blocks[1]["content"], "后续反思内容。")

    def test_merged_table_cells_are_not_duplicated(self):
        def build(document):
            document.add_paragraph("设计思路", style="Heading 2")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).merge(table.cell(0, 1)).text = "表格中的唯一设计说明。"
            document.add_paragraph("4.2 结果", style="Heading 2")

        blocks = self._parse(build)

        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0]["source_kind"], "table_cell")
        self.assertEqual(blocks[0]["content"], "表格中的唯一设计说明。")

    def test_fallback_is_explicit_and_keeps_units_separate(self):
        def build(document):
            document.add_paragraph("这是一段没有规范章节标题的实验报告正文，用于测试保守回退。")
            document.add_paragraph("另一段独立的正文内容，也应保持独立来源。")

        blocks = self._parse(build)

        self.assertEqual(len(blocks), 2)
        self.assertTrue(all(block["fallback"] for block in blocks))
        self.assertTrue(all(block["section_type"] == "GENERAL" for block in blocks))
        self.assertNotEqual(blocks[0]["source_index"], blocks[1]["source_index"])


if __name__ == "__main__":
    unittest.main()
